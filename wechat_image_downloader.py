import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import requests
from bs4 import BeautifulSoup
from PIL import Image, ImageTk
import io
import os
import tempfile
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class WeChatImageDownloader:
    def __init__(self, root):
        self.root = root
        self.root.title("微信公众号图片下载与Word排版工具")
        self.root.geometry("900x700")
        self.root.resizable(True, True)
        
        # 设置字体以支持中文
        self.font = ('SimHei', 10)
        
        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # URL输入区域
        self.url_frame = ttk.LabelFrame(self.main_frame, text="输入文章链接", padding="10")
        self.url_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.url_entry = ttk.Entry(self.url_frame, width=80, font=self.font)
        self.url_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        self.url_entry.insert(0, "https://mp.weixin.qq.com/s/")
        
        self.fetch_button = ttk.Button(self.url_frame, text="获取图片", command=self.fetch_images)
        self.fetch_button.pack(side=tk.RIGHT)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, pady=(0, 10))
        
        # 图片预览区域
        self.preview_frame = ttk.LabelFrame(self.main_frame, text="图片预览", padding="10")
        self.preview_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建画布用于显示图片
        self.canvas_frame = ttk.Frame(self.preview_frame)
        self.canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(self.canvas_frame, bg="white")
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 滚动条
        self.vscrollbar = ttk.Scrollbar(self.canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        self.vscrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.hscrollbar = ttk.Scrollbar(self.main_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
        self.hscrollbar.pack(fill=tk.X)
        
        self.canvas.config(yscrollcommand=self.vscrollbar.set, xscrollcommand=self.hscrollbar.set)
        
        # 创建图片容器框架
        self.images_frame = ttk.Frame(self.canvas)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.images_frame, anchor="nw")
        
        # 绑定事件以更新滚动区域
        self.images_frame.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)
        
        # 底部按钮区域
        self.button_frame = ttk.Frame(self.main_frame, padding="10")
        self.button_frame.pack(fill=tk.X, pady=(10, 0))
        
        self.select_all_button = ttk.Button(self.button_frame, text="全选", command=self.select_all)
        self.select_all_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.deselect_all_button = ttk.Button(self.button_frame, text="取消全选", command=self.deselect_all)
        self.deselect_all_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.save_button = ttk.Button(self.button_frame, text="保存选中图片", command=self.save_selected_images)
        self.save_button.pack(side=tk.LEFT, padx=(0, 10))
        
        # 修改为Word排版按钮
        self.word_button = ttk.Button(self.button_frame, text="生成Word文档", command=self.generate_word_document)
        self.word_button.pack(side=tk.LEFT)
        
        # 状态条
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        self.status_bar = ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 存储图片数据
        self.images_data = []  # 存储(Image对象, 临时路径, 复选框var)元组
        self.selected_count = 0
        
        # 设置请求头以模拟浏览器
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.8,zh-TW;q=0.7,zh-HK;q=0.5,en-US;q=0.3,en;q=0.2',
        }
        
    def on_frame_configure(self, event=None):
        """更新画布的滚动区域"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        
    def on_canvas_configure(self, event=None):
        """调整图片容器框架的宽度以适应画布"""
        width = event.width if event else self.canvas.winfo_width()
        self.canvas.itemconfig(self.canvas_window, width=width)
        
    def set_status(self, message):
        """更新状态条消息"""
        self.status_var.set(message)
        self.root.update_idletasks()
        
    def fetch_images(self):
        """从输入的URL中获取图片"""
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showerror("错误", "请输入有效的URL")
            return
        
        # 清空之前的图片
        for widget in self.images_frame.winfo_children():
            widget.destroy()
        self.images_data.clear()
        self.selected_count = 0
        
        self.set_status("正在获取文章内容...")
        self.progress_var.set(10)
        
        try:
            # 发送请求获取网页内容
            response = requests.get(url, headers=self.headers, timeout=10)
            response.raise_for_status()
            
            self.progress_var.set(30)
            self.set_status("正在解析网页内容...")
            
            # 解析HTML获取图片链接
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 微信公众号文章中图片的常见模式
            img_tags = soup.find_all('img')
            
            if not img_tags:
                messagebox.showinfo("提示", "未找到图片")
                self.set_status("就绪")
                self.progress_var.set(0)
                return
            
            # 提取有效图片链接
            image_urls = []
            for img in img_tags:
                # 安全地获取图片链接，避免索引越界错误
                img_url = None
                # 尝试从不同属性获取图片链接
                src = img.get('src', '')
                data_src = img.get('data-src', '')
                data_srcset = img.get('data-srcset', '')
                
                # 按优先级选择有效的图片链接
                if src and src.strip():
                    img_url = src
                elif data_src and data_src.strip():
                    img_url = data_src
                elif data_srcset and data_srcset.strip():
                    # 安全地处理data-srcset属性
                    try:
                        parts = data_srcset.split(',')[0].strip()
                        if parts:
                            img_url = parts.split()[0]  # 获取URL部分（忽略分辨率信息）
                    except IndexError:
                        # 如果解析失败，跳过此属性
                        pass
                
                # 验证图片链接并添加到列表
                if img_url and (img_url.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')) or 'mmbiz' in img_url):
                    # 补全相对URL
                    if not img_url.startswith(('http://', 'https://')):
                        parsed_url = urlparse(url)
                        base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
                        img_url = base_url + img_url if img_url.startswith('/') else base_url + '/' + img_url
                    image_urls.append(img_url)
            
            if not image_urls:
                messagebox.showinfo("提示", "未找到有效的图片链接")
                self.set_status("就绪")
                self.progress_var.set(0)
                return
            
            self.progress_var.set(40)
            self.set_status(f"找到{len(image_urls)}张图片，正在下载...")
            
            # 创建临时目录存储图片
            self.temp_dir = tempfile.mkdtemp()
            
            # 下载图片并显示预览
            for i, img_url in enumerate(image_urls):
                try:
                    # 下载图片
                    img_response = requests.get(img_url, headers=self.headers, timeout=10)
                    img_response.raise_for_status()
                    
                    # 打开图片
                    img_data = img_response.content
                    image = Image.open(io.BytesIO(img_data))
                    
                    # 保存临时文件
                    temp_path = os.path.join(self.temp_dir, f"img_{i}.{image.format.lower()}")
                    with open(temp_path, 'wb') as f:
                        f.write(img_data)
                    
                    # 调整图片大小以适应预览
                    max_width = 800
                    if image.width > max_width:
                        ratio = max_width / image.width
                        new_height = int(image.height * ratio)
                        image = image.resize((max_width, new_height), Image.LANCZOS)
                    
                    # 创建图片框架
                    img_frame = ttk.Frame(self.images_frame, padding="5")
                    img_frame.pack(fill=tk.X, pady=5)
                    
                    # 添加复选框
                    var = tk.BooleanVar()
                    checkbox = ttk.Checkbutton(img_frame, variable=var, command=lambda v=var: self.update_selected_count(v))
                    checkbox.pack(side=tk.LEFT, padx=5, pady=5)
                    
                    # 显示图片
                    photo = ImageTk.PhotoImage(image)
                    img_label = ttk.Label(img_frame, image=photo)
                    img_label.image = photo  # 保持引用
                    img_label.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
                    
                    # 显示图片信息
                    info_label = ttk.Label(img_frame, text=f"图片 {i+1}: {image.width}x{image.height}", font=self.font)
                    info_label.pack(side=tk.BOTTOM, fill=tk.X, pady=(5, 0))
                    
                    # 存储图片数据
                    self.images_data.append((image, temp_path, var))
                    
                    # 更新进度
                    progress = 40 + (i + 1) / len(image_urls) * 60
                    self.progress_var.set(progress)
                    
                except Exception as e:
                    print(f"下载图片 {i+1} 失败: {str(e)}")
                    continue
            
            if not self.images_data:
                messagebox.showinfo("提示", "所有图片下载失败")
                self.set_status("就绪")
                self.progress_var.set(0)
                return
            
            self.set_status(f"成功加载 {len(self.images_data)} 张图片")
            self.progress_var.set(100)
            
        except requests.exceptions.RequestException as e:
            messagebox.showerror("网络错误", f"无法连接到服务器: {str(e)}")
            self.set_status("就绪")
            self.progress_var.set(0)
        except Exception as e:
            messagebox.showerror("错误", f"处理过程中发生错误: {str(e)}")
            self.set_status("就绪")
            self.progress_var.set(0)
            
    def update_selected_count(self, var):
        """更新选中的图片数量"""
        if var.get():
            self.selected_count += 1
        else:
            self.selected_count -= 1
        self.set_status(f"已选择 {self.selected_count} 张图片")
        
    def select_all(self):
        """全选图片"""
        for _, _, var in self.images_data:
            var.set(True)
        self.selected_count = len(self.images_data)
        self.set_status(f"已选择 {self.selected_count} 张图片")
        
    def deselect_all(self):
        """取消全选图片"""
        for _, _, var in self.images_data:
            var.set(False)
        self.selected_count = 0
        self.set_status(f"已选择 {self.selected_count} 张图片")
        
    def save_selected_images(self):
        """保存选中的图片"""
        if self.selected_count == 0:
            messagebox.showinfo("提示", "请先选择要保存的图片")
            return
        
        # 选择保存目录
        save_dir = filedialog.askdirectory(title="选择保存目录")
        if not save_dir:
            return
        
        self.set_status(f"正在保存 {self.selected_count} 张图片...")
        
        saved_count = 0
        for i, (image, temp_path, var) in enumerate(self.images_data):
            if var.get():
                try:
                    # 获取原始图片
                    with open(temp_path, 'rb') as f:
                        img_data = f.read()
                    
                    # 保存图片
                    save_path = os.path.join(save_dir, f"wechat_image_{i+1}.{image.format.lower()}")
                    with open(save_path, 'wb') as f:
                        f.write(img_data)
                    
                    saved_count += 1
                except Exception as e:
                    print(f"保存图片 {i+1} 失败: {str(e)}")
                    continue
        
        self.set_status(f"成功保存 {saved_count} 张图片到 {save_dir}")
        messagebox.showinfo("完成", f"成功保存 {saved_count} 张图片")
        
    def generate_word_document(self):
        """将选中的图片按照A4纸张大小排版到Word文档中"""
        if self.selected_count == 0:
            messagebox.showinfo("提示", "请先选择要排版的图片")
            return
        
        # 让用户选择保存位置和文件名
        file_path = filedialog.asksaveasfilename(
            title="保存Word文档",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if not file_path:
            return
        
        self.set_status(f"正在生成Word文档...")
        
        try:
            # 获取选中的图片
            selected_images = [(image, temp_path, i) for i, (image, temp_path, var) in enumerate(self.images_data) if var.get()]
            
            if not selected_images:
                messagebox.showinfo("提示", "没有选中的图片")
                self.set_status("就绪")
                return
            
            # 创建Word文档
            doc = Document()
            
            # 设置页面为A4大小，边距为1厘米
            sections = doc.sections
            for section in sections:
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
            
            # 计算A4页面可用宽度和高度（使用固定值，避免类型问题）
            available_width_cm = 19  # A4宽度21cm减去左右边距各1cm
            available_height_cm = 27.7  # A4高度29.7cm减去上下边距各1cm
            
            # 添加图片到文档
            for i, (image, temp_path, original_index) in enumerate(selected_images):
                self.set_status(f"正在处理第 {i+1}/{len(selected_images)} 张图片...")
                
                # 为每张图片创建新段落
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 计算图片原始宽高比
                with Image.open(temp_path) as img:
                    original_width, original_height = img.size
                    aspect_ratio = original_height / original_width
                
                # 计算图片应该设置的宽度和高度，确保完全适应A4页面
                # 优先使用页面宽度
                target_width_cm = available_width_cm
                target_height_cm = target_width_cm * aspect_ratio
                
                # 如果按宽度缩放后高度超过页面高度，则按高度缩放
                if target_height_cm > available_height_cm:
                    target_height_cm = available_height_cm
                    target_width_cm = target_height_cm / aspect_ratio
                
                # 添加图片，使用计算出的宽度和高度
                run = paragraph.add_run()
                run.add_picture(temp_path, width=Cm(target_width_cm), height=Cm(target_height_cm))
            
            # 保存文档
            doc.save(file_path)
            
            self.set_status(f"成功生成Word文档")
            messagebox.showinfo("完成", f"Word文档已保存到：{file_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成Word文档时发生错误: {str(e)}")
            self.set_status("就绪")

if __name__ == "__main__":
    root = tk.Tk()
    # 设置中文字体支持
    root.option_add("*Font", ("SimHei", 10))
    app = WeChatImageDownloader(root)
    root.mainloop()